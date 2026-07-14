#!/usr/bin/env python3
"""
Build Word documents for the MAS IBM Bob demo framework.
Run:  python3 build_docs.py

Outputs (same directory as this script):
  MAS_Demo_Script.docx      — presenter guide for any customer demo
  CSE_How_To_Guide.docx     — reusable setup guide for all CSEs
"""

import os
import sys

# Install python-docx if needed
try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document

from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'output')
os.makedirs(HERE, exist_ok=True)


# ── helpers ──────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def para(doc, text, bold=False, italic=False, size=10, color=None, indent=0,
         sb=0, sa=4):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def bullet(doc, text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.4 + indent * 0.4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text).font.size = Pt(10)
    return p

def numbered(doc, text, indent=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.4 + indent * 0.4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text).font.size = Pt(10)
    return p

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    _shade(p, 'F0F0F0')
    return p

def callout(doc, label, text, fill='FFFDE7'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(f'{label}:  ')
    r.bold = True; r.font.size = Pt(10)
    p.add_run(text).font.size = Pt(10)
    _shade(p, fill)
    return p

def _shade(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True; run.font.size = Pt(9)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), 'BBBBBB')
    pBdr.append(bot); pPr.append(pBdr)

def margins(doc, t=2.5, b=2.5, l=3.0, r=3.0):
    s = doc.sections[0]
    s.top_margin = Cm(t); s.bottom_margin = Cm(b)
    s.left_margin = Cm(l); s.right_margin  = Cm(r)


# ── DOCUMENT 1: Demo Script ───────────────────────────────────────────────────

# ── NOAA demo data (loaded into Maximo BEDFORD site) ─────────────────────────
ASSETS = [
    ('NOAA-BUY-001', 'Pacific Ocean Monitoring Buoy - Station Alpha',      'SRN-2024-001'),
    ('NOAA-BUY-002', 'Gulf of Mexico Data Buoy - Station Bravo',           'SRN-2024-002'),
    ('NOAA-BUY-003', 'Atlantic Hurricane Tracking Buoy - Station Charlie', 'SRN-2024-003'),
    ('NOAA-BUY-004', 'Bering Sea Ice Monitoring Buoy - Station Delta',     'SRN-2024-004'),
    ('NOAA-BUY-005', 'Great Lakes Water Quality Buoy - Station Echo',      'SRN-2024-005'),
    ('NOAA-BUY-006', 'Chesapeake Bay Tidal Buoy - Station Foxtrot',        'SRN-2024-006'),
    ('NOAA-BUY-007', 'Caribbean Sea Wave Height Buoy - Station Golf',      'SRN-2024-007'),
    ('NOAA-BUY-008', 'Arctic Circle Drift Buoy - Station Hotel',           'SRN-2024-008'),
    ('NOAA-BUY-009', 'Coral Sea Reef Monitoring Buoy - Station India',     'SRN-2024-009'),
    ('NOAA-BUY-010', 'East Coast Storm Surge Buoy - Station Juliet',       'SRN-2024-010'),
]

WORK_ORDERS = [
    # (wonum, description, type, asset, lead_story)
    ('NOAA-WO-001', 'Annual Sensor Calibration - Pacific Buoy Alpha',           'PM', 'NOAA-BUY-001', False),
    ('NOAA-WO-002', 'Emergency Battery Pack Replacement - Gulf Buoy Bravo',     'CM', 'NOAA-BUY-002', True),   # ★ PRIMARY CM / high-drama
    ('NOAA-WO-003', 'Mooring Chain Inspection and Re-tensioning - Atlantic Buoy Charlie', 'PM', 'NOAA-BUY-003', False),
    ('NOAA-WO-004', 'Ice Sensor Array Replacement - Bering Sea Buoy Delta',     'CM', 'NOAA-BUY-004', False),
    ('NOAA-WO-005', 'Water Quality Sensor Suite Cleaning - Great Lakes Buoy Echo',  'PM', 'NOAA-BUY-005', False),
    ('NOAA-WO-006', 'Tidal Pressure Gauge Recertification - Chesapeake Buoy Foxtrot','PM', 'NOAA-BUY-006', False),
    ('NOAA-WO-007', 'Wave Rider Sensor Realignment - Caribbean Buoy Golf',      'CM', 'NOAA-BUY-007', False),
    ('NOAA-WO-008', 'Iridium Satellite Modem Replacement - Arctic Buoy Hotel',  'CM', 'NOAA-BUY-008', False),
    ('NOAA-WO-009', 'Coral Bleaching Sensor Upgrade - Reef Buoy India',         'PM', 'NOAA-BUY-009', False),
    ('NOAA-WO-010', 'Storm Surge Sensor Annual Overhaul - East Coast Buoy Juliet','PM','NOAA-BUY-010', False),
]

# Tasks and materials for the two featured WOs
WO002_TASKS = [
    'Task 1 — Confirm GPS position fix and isolate vessel approach vector',
    'Task 2 — Power-down buoy electronics via shore-side remote command',
    'Task 3 — Remove battery housing cover and photograph condition',
    'Task 4 — Extract failed lithium-ion battery pack (P/N BATT-LI-48V)',
    'Task 5 — Install replacement battery pack and torque housing bolts to 35 Nm',
    'Task 6 — Power-on sequence, verify all sensors report nominal readings',
]
WO002_MATERIALS = [
    ('BATT-LI-48V',   'Lithium-Ion Battery Pack 48V/200Ah',  1,   4200.00),
    ('SEAL-KIT-04',   'Weatherproof Gasket & Seal Kit',       1,    185.00),
    ('BOLT-M12-SS',   'Stainless Steel M12 Bolts (pack/10)', 2,     42.00),
    ('DIAG-HARNESS',  'Battery Diagnostic Harness',           1,    310.00),
]

WO001_TASKS = [
    'Task 1 — Download sensor calibration log via satellite link',
    'Task 2 — Compare pressure, temperature, and salinity readings against NIST standards',
    'Task 3 — Apply offset corrections using calibration software v4.2',
    'Task 4 — Run 24-hour validation data cycle and capture CSV export',
    'Task 5 — Update calibration certificate and log in Maximo work order',
    'Task 6 — Confirm next calibration due date — set 12-month PM trigger',
]
WO001_MATERIALS = [
    ('CAL-REF-STD',   'NIST Traceable Calibration Reference Standard', 1,  890.00),
    ('CABLE-USBC-IP', 'IP67-rated USB-C Sensor Interface Cable',        1,  130.00),
    ('LABEL-CERT',    'Calibration Certificate Label Set',              1,   22.00),
]


def build_demo_script():
    doc = Document()
    margins(doc)

    # Cover
    t = doc.add_heading('IBM Maximo Application Suite', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('Demo Script — Presenter Guide  |  NOAA Buoy Fleet Use Case')
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x3b, 0x82, 0xd4)

    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run(
        'Asset Management  ·  Work Management  ·  Inventory & Parts\n'
        'Audience: NOAA — National Oceanic and Atmospheric Administration    |    Runtime: 45–60 min\n'
        'Environment: MAS Manage  |  Site: BEDFORD  |  Org: EAGLENA'
    ).font.size = Pt(9)

    hr(doc); doc.add_paragraph()

    # ── Demo Data Reference ────────────────────────────────────────────────────
    set_heading(doc, 'Demo Data Reference  (loaded in Maximo)', 1, (0x1f, 0x23, 0x28))
    para(doc,
        'All records below are live in the Maximo BEDFORD site. Use these exact numbers '
        'when navigating during the demo — do not search generically.',
        italic=True, size=9)

    set_heading(doc, 'Assets  (NOAA-BUY-001 → NOAA-BUY-010, Site: BEDFORD)', 2)
    table(doc,
        ['Asset #', 'Description', 'Serial #', 'Location'],
        [(a, d, s, 'SHIPPING') for a, d, s in ASSETS],
        widths=[3.0, 8.5, 2.8, 2.2]
    )

    set_heading(doc, 'Work Orders  (NOAA-WO-001 → NOAA-WO-010, Status: WAPPR)', 2)
    table(doc,
        ['WO #', 'Description', 'Type', 'Asset'],
        [(w, d, t, a) for w, d, t, a, _ in WORK_ORDERS],
        widths=[2.8, 9.5, 1.2, 3.0]
    )

    callout(doc, 'Primary CM (Emergency)',
        'NOAA-WO-002 — Emergency Battery Pack Replacement - Gulf Buoy Bravo\n'
        'Asset: NOAA-BUY-002 | Serial: SRN-2024-002 | Work Type: CM\n'
        'This is your high-drama opening: battery failure on an active hurricane-season monitoring buoy.',
        'FFF0F0')

    callout(doc, 'Primary PM (Preventive)',
        'NOAA-WO-001 — Annual Sensor Calibration - Pacific Buoy Alpha\n'
        'Asset: NOAA-BUY-001 | Serial: SRN-2024-001 | Work Type: PM\n'
        'Use this to contrast reactive vs. proactive maintenance in Scene 2.',
        'F0FDF4')

    set_heading(doc, 'NOAA-WO-002 Task Steps (Plans tab → Tasks)', 2)
    for t_text in WO002_TASKS:
        bullet(doc, t_text)

    set_heading(doc, 'NOAA-WO-002 Planned Materials (Plans tab → Materials)', 2)
    table(doc,
        ['Part #', 'Description', 'Qty', 'Unit Cost'],
        [(p, d, q, f'${c:,.2f}') for p, d, q, c in WO002_MATERIALS],
        widths=[2.8, 7.5, 1.2, 2.0]
    )

    set_heading(doc, 'NOAA-WO-001 Task Steps (Plans tab → Tasks)', 2)
    for t_text in WO001_TASKS:
        bullet(doc, t_text)

    set_heading(doc, 'NOAA-WO-001 Planned Materials (Plans tab → Materials)', 2)
    table(doc,
        ['Part #', 'Description', 'Qty', 'Unit Cost'],
        [(p, d, q, f'${c:,.2f}') for p, d, q, c in WO001_MATERIALS],
        widths=[2.8, 7.5, 1.2, 2.0]
    )

    doc.add_paragraph(); hr(doc)

    # ── Pre-Demo Setup ────────────────────────────────────────────────────────
    set_heading(doc, 'Pre-Demo Setup', 1, (0x1f, 0x23, 0x28))
    para(doc,
        'Data is already loaded — no script run needed for this NOAA demo. '
        'Complete the browser setup below before the audience joins.',
        italic=True, size=10)

    set_heading(doc, 'Pre-Demo Browser Setup', 2)
    para(doc, 'Pre-open these exact tabs before sharing your screen:', bold=True, size=10)
    for item in [
        'Tab 1 — Assets → Assets → filter Site = BEDFORD, Asset # starts with NOAA-BUY',
        'Tab 2 — Asset NOAA-BUY-002 (Gulf Buoy Bravo) — your emergency/CM opening record',
        'Tab 3 — Work Orders → Work Order Tracking → Site = BEDFORD, Status = WAPPR',
        'Tab 4 — Work Order NOAA-WO-002 (Emergency Battery Pack Replacement) — Plans tab open',
        'Tab 5 — Work Order NOAA-WO-001 (Annual Sensor Calibration) — for PM contrast',
        'Tab 6 — Inventory → Inventory → Storeroom = BEDFORD (for Scene 3)',
    ]:
        bullet(doc, item)
    bullet(doc, 'Close all unrelated applications and browser tabs before sharing your screen.')
    doc.add_paragraph()

    hr(doc)

    # ── Opening Hook ──────────────────────────────────────────────────────────
    set_heading(doc, 'Opening Hook', 1, (0x7c, 0x5c, 0xd8))
    para(doc, 'Deliver this before switching to Maximo. No slides required.', italic=True, size=9)

    callout(doc, 'Say this',
        '"NOAA operates hundreds of ocean monitoring buoys across every major body of water '
        'in the Western Hemisphere. These buoys feed real-time data to hurricane track models, '
        'tsunami warning systems, and climate research. When one goes offline — even temporarily — '
        'there is a gap in the data that covers thousands of square miles of ocean.\n\n'
        'Yesterday, Station Bravo in the Gulf of Mexico — buoy NOAA-BUY-002 — triggered a '
        'battery failure alert. The sensor suite is still on backup power. You have a narrow '
        'maintenance window before it goes dark during peak Atlantic hurricane season.\n\n'
        'Your maintenance team needs to answer right now: Who goes? What parts? '
        'How long will this take? Who authorizes the deployment?\n\n'
        'That is exactly what we\'re going to walk through. Everything on screen is live — '
        'this is your data, in your system."',
        'F5F0FF')

    para(doc, 'Transition: "Let\'s start with the asset — buoy NOAA-BUY-002."', italic=True, size=9)
    doc.add_paragraph()

    # ── Scene 1: Asset Management ─────────────────────────────────────────────
    set_heading(doc, 'Scene 1 — Asset Management  (10 min)', 1, (0x3b, 0x82, 0xd4))
    para(doc, 'Navigate to:  Tab 1 (Assets list, BEDFORD) → open NOAA-BUY-002', bold=True, size=10)

    set_heading(doc, 'What to Show', 2)
    bullet(doc, 'Asset list with all 10 NOAA-BUY-001 → 010 buoys visible — fleet-level view.')
    bullet(doc, 'Open NOAA-BUY-002 — "Gulf of Mexico Data Buoy - Station Bravo", Serial: SRN-2024-002, Location: SHIPPING.')
    bullet(doc, 'Work Order History tab — NOAA-WO-002 is already linked to this asset.')
    bullet(doc, 'Asset header fields: description, serial number, location SHIPPING, site BEDFORD.')
    bullet(doc, 'Then briefly open NOAA-BUY-001 to show a healthy, PM-maintained counterpart.')

    set_heading(doc, 'Narration Guide', 2)
    callout(doc, 'Core message',
        '"Every buoy in the NOAA fleet has a single record in Maximo. Serial number SRN-2024-002 '
        'is here. The complete maintenance history — every repair, every calibration — is here. '
        'Nothing lives in a spreadsheet or a technician\'s notebook. '
        'When Station Bravo fails, we have the full picture in seconds."',
        'EFF6FF')

    set_heading(doc, 'Key Talking Points', 2)
    bullet(doc, 'All 10 buoys (Alpha through Juliet) are in one system. No cross-referencing databases by region.')
    bullet(doc, 'Serial number SRN-2024-002 ties this physical unit to its entire service history — critical for warranty and failure-pattern analysis.')
    bullet(doc, 'Location SHIPPING shows where the buoy is staged — drives logistics for who deploys and which storeroom is nearest.')
    bullet(doc, 'Anticipated question: "We have hundreds more buoys not in the demo." → "Maximo has bulk import tools. You start with your highest-criticality assets and migrate the rest incrementally."')
    para(doc, 'Transition: "The battery alert triggered a work order. Let\'s look at that next."', italic=True, size=9)
    doc.add_paragraph()

    # ── Scene 2: Work Order Management ───────────────────────────────────────
    set_heading(doc, 'Scene 2 — Work Order Management  (20 min)', 1, (0x05, 0x96, 0x69))
    para(doc, 'Navigate to:  Tab 3 (WO list, BEDFORD, WAPPR) → open NOAA-WO-002', bold=True, size=10)

    set_heading(doc, 'Step 1 — Show the Work Order Queue', 2)
    bullet(doc, 'Work order list shows all 10 NOAA-WO-001 → 010 at Status WAPPR.')
    bullet(doc, 'Point out the type mix: CM (WO-002, 004, 007, 008 — corrective failures) vs PM (WO-001, 003, 005, 006, 009, 010 — scheduled).')
    bullet(doc, 'NOAA-WO-002 is your high-drama record — open it.')

    set_heading(doc, 'Step 2 — Open NOAA-WO-002 (Emergency CM)', 2)
    callout(doc, 'Core message',
        '"Work order NOAA-WO-002: Emergency Battery Pack Replacement on Gulf Buoy Bravo. '
        'Corrective Maintenance. This WO was created the moment the battery alert fired. '
        'One supervisor click away from triggering crew deployment and parts reservation."',
        'F0FDF4')
    bullet(doc, 'WO header: Description = "Emergency Battery Pack Replacement - Gulf Buoy Bravo", Asset = NOAA-BUY-002, Type = CM, Status = WAPPR.')
    bullet(doc, 'Plans tab → Tasks: show the 6 task steps (confirm GPS fix → power-down → photograph → extract battery → install replacement → power-on validation).')
    bullet(doc, 'Plans tab → Materials: 4 planned parts — BATT-LI-48V ($4,200), SEAL-KIT-04 ($185), BOLT-M12-SS ($42 × 2), DIAG-HARNESS ($310). Maximo calculates estimated material cost automatically.')
    bullet(doc, 'Total estimated material cost for this WO: ~$4,779.')

    set_heading(doc, 'Step 3 — Approve the Work Order', 2)
    bullet(doc, 'Change status: WAPPR → APPR.')
    bullet(doc, 'Explain: this single action triggers parts reservation against the storeroom, notifies the assigned crew, and makes NOAA-WO-002 visible to every stakeholder in real time.')
    bullet(doc, 'Revert status back to WAPPR after demonstrating (to keep demo environment clean).')

    set_heading(doc, 'Step 4 — Switch to NOAA-WO-001 (PM Contrast)', 2)
    bullet(doc, 'Navigate to Tab 5 — open NOAA-WO-001: "Annual Sensor Calibration - Pacific Buoy Alpha".')
    bullet(doc, 'Type = PM. Tasks: 6 calibration steps (download log → compare to NIST standards → apply offset corrections → 24hr validation run → update certificate → set next PM trigger).')
    bullet(doc, 'Materials: CAL-REF-STD ($890), CABLE-USBC-IP ($130), LABEL-CERT ($22).')
    callout(doc, 'Say this',
        '"That WO-002 was reactive — the battery failed, the buoy alarmed, and we responded. '
        'This WO-001 is proactive — we know Station Alpha\'s sensors drift after 12 months, '
        'so the calibration is scheduled and triggered automatically. '
        'The cost of a planned calibration is a fraction of an emergency sensor replacement."',
        'F0FDF4')

    set_heading(doc, 'Key Talking Points', 2)
    bullet(doc, '"Every maintenance event on every NOAA buoy is traceable. Maximo timestamps every action and attributes it to a technician — that is your chain of custody."')
    bullet(doc, '"A work order is not a ticket — it is the legal, operational, and financial record of the work. It connects the asset failure, the parts consumed, the labor hours, and the resolution."')
    bullet(doc, '"PM schedules — like WO-001 — mean fewer emergency deployments like WO-002. Each emergency response costs 3–5× more than a planned maintenance visit."')
    bullet(doc, 'Anticipated question: "How does this connect to our satellite monitoring system?" → "Maximo\'s Integration Framework can receive telemetry alerts via REST webhook and auto-create pre-populated work orders — eliminating manual triage entirely."')
    para(doc, 'Transition: "Let\'s check whether the battery pack is in stock."', italic=True, size=9)
    doc.add_paragraph()

    # ── Scene 3: Inventory ────────────────────────────────────────────────────
    set_heading(doc, 'Scene 3 — Inventory & Parts Management  (10 min)', 1, (0xd9, 0x77, 0x06))
    para(doc, 'Navigate to:  Tab 6 — Inventory → Inventory → Storeroom BEDFORD', bold=True, size=10)

    set_heading(doc, 'What to Show', 2)
    bullet(doc, 'Storeroom view — all items, quantities on hand, reorder points.')
    bullet(doc, 'Reference BATT-LI-48V (Lithium-Ion Battery Pack 48V/200Ah) — the $4,200 part on NOAA-WO-002. Show quantity on hand and reorder point.')
    bullet(doc, 'Reorder point concept: when stock drops to 1 unit, Maximo auto-triggers a Purchase Requisition — so the next emergency deployment is never delayed by a stockout.')
    bullet(doc, 'Show other WO-002 materials: SEAL-KIT-04, BOLT-M12-SS, DIAG-HARNESS — all planned and costed on the work order before the technician ships out.')

    set_heading(doc, 'Key Talking Points', 2)
    bullet(doc, '"The difference between a 2-day repair and a 3-week repair on a remote ocean buoy is often one part. BATT-LI-48V is $4,200. Not having it on the shelf costs far more in missed data coverage."')
    bullet(doc, '"Maximo\'s material planning lets your team see 6 months of upcoming work orders and identify parts gaps before they become emergency procurement events."')
    bullet(doc, '"Every purchase requisition is traceable to the asset and work order it supports — that is your audit trail from budget to wrench."')
    bullet(doc, 'Anticipated question: "Does this connect to our ERP / procurement system?" → "Yes — Maximo has certified integrations with SAP and Oracle, plus generic REST connectors for any purchasing system."')
    doc.add_paragraph()

    # ── Scene 4: Reporting ────────────────────────────────────────────────────
    set_heading(doc, 'Scene 4 — Reporting & Management Visibility  (6 min)', 1, (0xdc, 0x26, 0x26))

    set_heading(doc, 'What to Show', 2)
    bullet(doc, 'Work order list — breakdown: 4 CM (corrective/emergency) vs 6 PM (scheduled) across the buoy fleet.')
    bullet(doc, 'PM compliance: of the 6 scheduled WOs, how many are approved and on schedule vs overdue.')
    bullet(doc, 'Cost tracking: estimated material cost across all 10 WOs — highlight NOAA-WO-002 as the highest-cost emergency event.')
    bullet(doc, 'Optional: Maximo Predict / AI — asset health scores for each buoy, failure probability by station, recommended PM windows before hurricane season.')

    callout(doc, 'Closing message',
        '"NOAA leadership needs answers to: How many buoys are at risk right now? '
        'What is our PM completion rate across the fleet? '
        'What did we spend on emergency repairs vs planned maintenance this quarter? '
        'Maximo delivers those answers in real time — no manual reports, no data exports."',
        'FEF2F2')
    doc.add_paragraph()

    # ── Call to Action ────────────────────────────────────────────────────────
    set_heading(doc, 'Closing & Call to Action  (3 min)', 1, (0x1f, 0x23, 0x28))
    callout(doc, 'Say this',
        '"What you\'ve seen today is the complete operational loop for a single buoy failure: '
        'Station Bravo alerted, a work order was created, the right parts were identified, '
        'the work was approved, and leadership had full visibility — in minutes, not days.\n\n'
        'Scale that across hundreds of buoys, in every ocean, with automated PM triggers and '
        'predictive health scores, and your team shifts from responding to emergencies '
        'to preventing them.\n\n'
        'Our recommended next step is a two-day discovery workshop with your operations '
        'and IT teams. We map your current processes, your existing fleet data, and your '
        'workflows — and we produce a concrete MAS implementation roadmap with your numbers. '
        'Does that sound like a worthwhile next step?"',
        'F0F0F0')
    doc.add_paragraph()

    # ── Objection Handling ────────────────────────────────────────────────────
    set_heading(doc, 'Objection Handling', 1, (0x1f, 0x23, 0x28))
    objections = [
        ('"We already track this in spreadsheets / legacy systems."',
         'Maximo is not a rip-and-replace on day one. Start with a pilot asset class — '
         'for NOAA, that could be your highest-criticality deepwater buoys. Prove the value, '
         'then migrate the rest of the fleet incrementally. Legacy systems run in parallel.'),
        ('"Procurement / acquisition takes too long."',
         'IBM has GSA Schedule 70 and NASA SEWP contract vehicles that compress full competition '
         'timelines from years to months. We can connect you with our federal contracts team.'),
        ('"We need FedRAMP authorization."',
         'MAS is available on IBM Cloud for Government (FedRAMP Moderate Authorized). '
         'We can provide the current authorization package on request.'),
        ('"We are worried about vendor lock-in."',
         'Maximo stores all data in open SQL databases with published REST APIs. Your data is '
         'yours and exportable at any time. There is also a large ecosystem of certified '
         'implementation partners beyond IBM.'),
        ('"What is the total cost of ownership?"',
         'The ROI case is built around three numbers: reduction in emergency deployments, '
         'reduction in unplanned data outages, and labor hours recaptured from manual tracking. '
         'We quantify those with your actual fleet data during the discovery workshop.'),
    ]
    for q, a in objections:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(q)
        r.bold = True; r.font.size = Pt(10)
        para(doc, a, indent=0.5, size=10, sa=6)

    doc.add_paragraph()

    # ── Q&A ───────────────────────────────────────────────────────────────────
    set_heading(doc, 'Anticipated Q&A', 1, (0x1f, 0x23, 0x28))
    qa = [
        ('Can Maximo receive alerts directly from our buoy telemetry / satellite feed?',
         'Yes. Maximo\'s Integration Framework and REST API can receive webhook alerts from any '
         'telemetry platform — satellite, GOES, NDC — and automatically create a pre-populated '
         'work order, eliminating manual triage. You can also close the loop by pushing '
         'WO completion status back to the monitoring system.'),
        ('Does Maximo work for field technicians deploying on ships or in remote locations?',
         'Maximo Mobile is an offline-first app. Technicians download WO-002 before boarding '
         'the deployment vessel, complete all 6 task steps and capture photos offline at sea, '
         'then sync when connectivity is restored at port.'),
        ('How long does implementation take for a fleet this size?',
         'A focused implementation for the buoy asset class typically runs 12–16 weeks. '
         'IBM has deployment accelerators for federal science agencies that compress this '
         'timeline, including pre-built data migration templates.'),
        ('What data does Maximo Predict need to generate health scores for buoys?',
         'It starts with historical work order data already in Maximo — failure codes, '
         'CM-to-PM ratios, mean time between failures by station. Over time you add '
         'real-time sensor feeds (battery voltage, sensor drift metrics). '
         'The more data, the more precise the predicted failure window.'),
        ('Can Maximo manage the entire asset lifecycle, including decommissioning?',
         'Yes — asset status tracks from Active through Decommissioned. Maximo captures '
         'disposal records, regulatory compliance dates, and can trigger replacement '
         'requisitions automatically when an asset reaches end-of-life threshold.'),
    ]
    for q, a in qa:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(2)
        p.add_run(f'Q:  {q}').bold = True
        para(doc, f'A:  {a}', indent=0.5, size=10, sa=6)

    doc.add_paragraph(); hr(doc)

    # ── Pre-Demo Checklist ────────────────────────────────────────────────────
    set_heading(doc, 'Pre-Demo Checklist', 1, (0x1f, 0x23, 0x28))
    for item in [
        'Logged into Maximo as the demo user (site: BEDFORD)',
        'Confirmed all 10 NOAA-BUY-001 → 010 assets exist in Assets → Assets',
        'Confirmed all 10 NOAA-WO-001 → 010 work orders exist at Status WAPPR',
        'Opened NOAA-BUY-002 and NOAA-WO-002 — Plans tab showing tasks and materials',
        'Tab 1: Assets list filtered to BEDFORD / NOAA-BUY',
        'Tab 2: Asset NOAA-BUY-002 open',
        'Tab 3: WO list filtered to BEDFORD, Status WAPPR',
        'Tab 4: Work Order NOAA-WO-002 open, Plans tab visible',
        'Tab 5: Work Order NOAA-WO-001 open',
        'Tab 6: Inventory → storeroom BEDFORD',
        'All unrelated applications and browser tabs closed',
        'Screen share tested before audience joins',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.add_run('☐  ')
        p.add_run(item).font.size = Pt(10)

    doc.add_paragraph()
    para(doc, 'IBM Maximo Application Suite — Customer Success Engineering  |  NOAA Demo', italic=True, size=8,
         color=(0x57, 0x60, 0x6a))

    out = os.path.join(HERE, 'MAS_Demo_Script.docx')
    doc.save(out)
    print(f'✓ {out}')
    return out


# ── DOCUMENT 2: CSE How-To Guide ─────────────────────────────────────────────

def build_how_to_guide():
    doc = Document()
    margins(doc)

    # Cover
    t = doc.add_heading('CSE How-To Guide', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('Bob + Maximo Manage  —  Demo Data Framework')
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x3b, 0x82, 0xd4)

    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run('Audience: IBM Customer Success Engineers  |  Setup Time: ~30 min per environment\n'
              'Requires: Bob, Python 3, MAS Manage API key').font.size = Pt(9)

    hr(doc)

    # Overview
    doc.add_paragraph()
    set_heading(doc, 'Overview', 1, (0x1f, 0x23, 0x28))
    para(doc,
        'This guide gives every CSE the exact steps to connect IBM Bob to any Maximo Application '
        'Suite (MAS) Manage instance, bulk-load industry-specific demo data, and generate a full '
        '1-hour demo script — all from within Bob. No external tooling is required beyond Python 3.',
        size=10)

    para(doc, 'What Bob can do once connected to Maximo:', bold=True, size=10)
    for b in [
        'Query live Maximo data — assets, work orders, inventory — in natural language',
        'Bulk-create demo assets, work orders, task operations, and planned materials via the Maximo REST API',
        'Generate a full 1-hour demo script with talking points, navigation steps, and Q&A prep',
        'Render results as formatted tables directly in the chat window',
    ]:
        bullet(doc, b)

    table(doc,
        ['File', 'Purpose'],
        [
            ['demo_loader.py',        'Template Python script — fill in config, run, all data is created in Maximo'],
            ['MAS_Demo_Script.docx',  'Presenter guide — opening hook, 4 demo scenes, objections, Q&A, checklist'],
            ['CSE_How_To_Guide.docx', 'This document — setup, querying, loading data, troubleshooting'],
            ['build_docs.py',         'Regenerates both .docx files from source — run after editing this script'],
        ],
        widths=[5.0, 9.5]
    )
    doc.add_paragraph()

    # ── Part 1 ────────────────────────────────────────────────────────────────
    set_heading(doc, 'Part 1: Connect Bob to Your Maximo Instance', 1, (0x3b, 0x82, 0xd4))

    set_heading(doc, 'Step 1 — Get Your Maximo URL and API Key', 2)
    table(doc,
        ['Item', 'Where to Find It'],
        [
            ['Maximo Manage URL',
             'TechZone reservation → Published Services → "Maximo Manage URL"\n'
             'Pattern:  https://<appid>.manage.<cluster>/maximo'],
            ['API Key',
             'Log into Maximo → top-right avatar → API Keys → Generate new key'],
        ],
        widths=[4.5, 10.0]
    )
    callout(doc, 'URL Gotcha',
        'The MAS Navigator / home URL (e.g. gtmdemo.home.inst1...) is NOT the API URL. '
        'The Manage API lives on the manage subdomain: gtmdemo.manage.inst1...\n'
        'Using the wrong URL causes 302 redirects to the SSO login page on every call.',
        'FFF8E1')

    set_heading(doc, 'Step 2 — Configure the maximo-mcp Server in Bob', 2)
    para(doc, 'Open  ~/.bob/settings/mcp.json  and merge in the entry below.', size=10)
    para(doc, 'NEVER overwrite the whole file — other servers will be deleted.', bold=True, size=10)
    code(doc, '''{
  "mcpServers": {
    "maximo-mcp": {
      "command": "npx",
      "args": ["-y", "maximo-mcp-server"],
      "env": {
        "MAXIMO_URL": "https://<manage-subdomain>/maximo",
        "MAXIMO_API_KEY": "<your-api-key>"
      },
      "alwaysAllow": [
        "list_object_structures", "get_schema_details",
        "query_maximo", "render_carbon_table",
        "render_carbon_details", "get_instance_details"
      ]
    }
  }
}''')
    callout(doc, 'Critical',
        'The env var MUST be  MAXIMO_URL  — NOT  MAXIMO_BASE_URL.\n'
        'The npm package silently ignores any other name, leaving instanceUrl empty.',
        'FFE4E4')

    set_heading(doc, 'Step 3 — Restart the MCP Server', 2)
    code(doc, 'Cmd+Shift+P  →  MCP: Restart Server  →  maximo-mcp')
    para(doc, 'Verify: ask Bob "Run get_instance_details for Maximo" — if instanceUrl is populated, you are connected.', size=10)

    set_heading(doc, 'Step 4 — Set the Default Insert Site in Maximo  ⚠ Most Common Failure Point', 2)
    callout(doc, 'Why this matters',
        'If the Maximo user account has querywithsite=true but no Default Insert Site configured, '
        'every POST will fail with "null is not a valid site" — even when siteid is in the request body. '
        'This must be done once per user account per environment.',
        'FFE4E4')
    para(doc, 'Fix (takes 60 seconds):', bold=True, size=10)
    numbered(doc, 'Log into Maximo as the API key user')
    numbered(doc, 'Click name/avatar → Profile  (or search "Default Information")')
    numbered(doc, 'Set  Default Insert Site  →  your demo site  (e.g. BEDFORD)')
    numbered(doc, 'Set  Storeroom Site  and  Default Storeroom  if demoing inventory')
    numbered(doc, 'Click OK — this persists permanently for the account')
    doc.add_paragraph()

    # ── Part 2 ────────────────────────────────────────────────────────────────
    set_heading(doc, 'Part 2: Query Maximo from Bob', 1, (0x3b, 0x82, 0xd4))
    para(doc, 'Once connected, ask Bob natural-language questions — it calls the Maximo REST API and renders results.', size=10)

    set_heading(doc, 'Example Prompts', 2)
    for p_text in [
        '"Show me 5 open work orders from Maximo"',
        '"Query mxapiasset for assets at site BEDFORD, show assetnum, description, status"',
        '"What object structures are available in this Maximo instance?"',
        '"List all valid locations at site BEDFORD"',
    ]:
        code(doc, p_text)

    set_heading(doc, 'Key Object Structures', 2)
    table(doc,
        ['Use Case', 'Object Structure'],
        [
            ['Work Orders',        'mxapiwodetail'],
            ['Assets',             'mxapiasset'],
            ['Inventory Items',    'mxapiitem'],
            ['Locations',          'mxapilocation'],
            ['Persons / Users',    'mxapiperson, mxapiperuser'],
            ['Companies/Vendors',  'mxapicompany'],
        ],
        widths=[5.5, 8.0]
    )
    doc.add_paragraph()

    # ── Part 3 ────────────────────────────────────────────────────────────────
    set_heading(doc, 'Part 3: Load Demo Data', 1, (0x7c, 0x5c, 0xd8))

    set_heading(doc, 'Option A — Ask Bob to Generate Everything', 2)
    callout(doc, 'Bob Prompt',
        '"I have a demo for [CUSTOMER] focused on [USE CASE / INDUSTRY]. '
        'Help me create custom demo assets and work orders for it. '
        'The site is [SITE]. Create 10 assets, 1 work order per asset with task steps '
        'and planned materials. Assign all work orders to [PERSON ID]."',
        'F3F0FF')
    para(doc, 'Bob will: query Maximo for valid locations and person IDs → generate industry-appropriate data → '
         'write the loader script → run it → confirm all records created.', italic=True, size=10)

    set_heading(doc, 'Option B — Adapt demo_loader.py Manually', 2)
    numbered(doc, 'Open demo_loader.py in this folder.')
    numbered(doc, 'Fill in the CONFIGURATION block at the top:')
    code(doc, '''LOGIN_BASE = "https://<frontend-host>/maximo"
API_BASE   = "https://<api-host>/maximo"
API_KEY    = "<your-api-key>"
SITE       = "<SITE>"       # e.g. BEDFORD
ORG        = "<ORG>"        # e.g. EAGLENA
LOCATION   = "<LOCATION>"  # e.g. CENTRAL
OWNER      = "<PERSON ID>" # PERSON.PERSONID — may contain spaces''')
    callout(doc, 'Person ID Gotcha',
        'OWNER must be the value of PERSON.PERSONID — not the login ID. '
        'These often differ (login: jsmith, personid: JOHN SMITH).\n'
        'Ask Bob: "Query mxapiperson for personid and displayname, 10 records"',
        'FFF8E1')
    numbered(doc, 'Replace the ASSETS list with your customer-specific records:')
    code(doc, '''{"assetnum": "CUST-AST-001",
 "description": "Asset description",
 "serialnum": "SN-2024-001",
 "location": LOCATION}    # must be a valid Maximo location''')
    callout(doc, 'Location Warning',
        'Location codes must exist in Maximo — never invent them. '
        'Ask Bob: "Query mxapilocation for location and description at site BEDFORD".\n'
        'Also omit: vendor, assettype, status — Maximo validates these against internal lists '
        'and will 400 if the value is not in the list.',
        'FFF8E1')
    numbered(doc, 'Replace the WORK_ORDERS list:')
    code(doc, '''{"wonum": "CUST-WO-001", "assetnum": "CUST-AST-001",
 "description": "Work order description",
 "worktype": "PM",          # valid: PM, CM, CP, EM
 "wopriority": 2,           # integer 1-5
 "tasks": ["Step 1", "Step 2", "Step 3"],
 "materials": [
   {"itemnum": "PART-001", "description": "Part name",
    "qty": 1, "unitcost": 100.00}
 ]}''')
    callout(doc, 'Materials Note',
        'The script uses linetype: MATERIAL (freeform). Parts do NOT need to pre-exist '
        'in the Maximo item master. This lets you demo against any fresh environment.',
        'F0FDF4')
    numbered(doc, 'Run:')
    code(doc, 'python3 demo_loader.py')
    para(doc, 'Assets: re-runnable (duplicate = skip). Work orders: NOT re-runnable without unique WO numbers. '
         'Delete existing WOs from the UI before re-running if needed.', italic=True, size=9)
    doc.add_paragraph()

    # ── Part 4 ────────────────────────────────────────────────────────────────
    set_heading(doc, 'Part 4: Generate the Demo Script via Bob', 1, (0x7c, 0x5c, 0xd8))
    callout(doc, 'Bob Prompt',
        '"Generate a 1-hour demo script for [CUSTOMER] covering asset management, '
        'work order management, inventory, and reporting. Include [INDUSTRY]-specific '
        'talking points, anticipated Q&A, and a pre-demo checklist."',
        'F3F0FF')
    para(doc, 'Bob produces a formatted HTML artifact with:', bold=True, size=10)
    table(doc,
        ['Section', 'Duration', 'Contents'],
        [
            ['Introduction',          '~8 min',  'Scene-setting narrative, platform overview'],
            ['Asset Management',      '~14 min', 'Fleet/asset view, record deep-dive, lifecycle history'],
            ['Work Order Management', '~20 min', 'CM emergency, PM scheduled, task steps, approval flow'],
            ['Inventory & Parts',     '~10 min', 'Planned materials, reorder points, cost coding'],
            ['Reporting',             '~6 min',  'PM compliance, cost visibility, AI/Predict preview'],
            ['Q&A',                   '~2 min',  'Pre-loaded answers: integration, mobile, FedRAMP, timeline'],
        ],
        widths=[4.0, 2.5, 8.0]
    )
    doc.add_paragraph()

    # ── Part 5: Troubleshooting ───────────────────────────────────────────────
    set_heading(doc, 'Part 5: Troubleshooting Reference', 1, (0xdc, 0x26, 0x26))

    set_heading(doc, 'Test URL and API Key Before Running Scripts', 2)
    code(doc, '''BASE="https://<your-manage-host>/maximo"
KEY="<your-api-key>"

# Login test — should return JSON with maxupg version
curl -s "$BASE/api/login?apikey=$KEY" -k

# Data query test — should return work order records
curl -s "$BASE/api/os/mxapiwodetail?lean=1&oslc.pageSize=3&oslc.select=wonum,status" \\
  -H "apikey: $KEY" -k | python3 -m json.tool''')

    set_heading(doc, 'Error Reference', 2)
    table(doc,
        ['Error', 'Root Cause', 'Fix'],
        [
            ['instanceUrl empty in Bob',         'Wrong env var name',               'MAXIMO_BASE_URL → MAXIMO_URL in mcp.json, restart server'],
            ['Failed to parse URL',              'MAXIMO_URL not set',               'Save mcp.json, restart MCP server'],
            ['null is not a valid site',         'No default insert site on user',   'Set Default Insert Site in Maximo UI — Part 1, Step 4'],
            ['Location X is not valid',          'Invented location code',           'Query mxapilocation for valid codes at your site'],
            ['Type OPERATING not in value list', 'Invalid assettype or status',      'Remove assettype and status from asset payload'],
            ['Must enter a valid company',       'vendor not in company master',     'Remove vendor from asset payload'],
            ['Person does not exist',            'Wrong owner — login ID vs personid','Query mxapiperson — use personid, not login ID'],
            ['Work Type MOD not valid',          'Non-standard work type',           'Valid values: PM, CM, CP, EM'],
            ['302 redirect on every call',       'Using Navigator URL not API URL',  'Use manage subdomain, not home/nav subdomain'],
            ['503 Application not available',    'MAS pod cold-starting',            'Wait 10–15 min — TechZone envs idle and need warm-up time'],
        ],
        widths=[4.5, 4.5, 5.5]
    )

    set_heading(doc, 'MAS URL Pattern', 2)
    code(doc, '''Navigator (UI):   https://<appid>.home.<cluster>/
API (Manage):     https://<appid>.manage.<cluster>/maximo
Data calls:       https://<appid>-all.manage.<cluster>/maximo  (redirect destination)''')
    para(doc, 'After login, Maximo may redirect to an *-all.manage.* hostname for data requests. '
         'The demo_loader.py script handles this by logging in on the manage host and calling data via the -all host.',
         italic=True, size=9)
    doc.add_paragraph()

    # ── Quick-Start Checklist ─────────────────────────────────────────────────
    hr(doc)
    set_heading(doc, 'Quick-Start Checklist', 1, (0x1f, 0x23, 0x28))
    for item in [
        'Maximo Manage URL identified — manage subdomain, not home/nav',
        'API key generated in Maximo for your demo user',
        'mcp.json updated with MAXIMO_URL (not MAXIMO_BASE_URL)',
        'MCP server restarted in Bob',
        'Bob confirms instanceUrl populated (get_instance_details)',
        'Default Insert Site set in Maximo UI user profile',
        'Valid location codes confirmed (query mxapilocation)',
        'Valid person ID confirmed (query mxapiperson)',
        'demo_loader.py configuration block filled in',
        'python3 demo_loader.py ran — all assets and WOs show ✓',
        'Demo data spot-checked in Maximo UI',
        'Demo script generated via Bob and reviewed',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run('☐  ')
        p.add_run(item).font.size = Pt(10)

    doc.add_paragraph()
    para(doc, 'IBM Maximo Application Suite — Customer Success Engineering', italic=True, size=8,
         color=(0x57, 0x60, 0x6a))

    out = os.path.join(HERE, 'CSE_How_To_Guide.docx')
    doc.save(out)
    print(f'✓ {out}')
    return out


if __name__ == '__main__':
    print('Building MAS IBM Bob demo documents...')
    print('=' * 50)
    p1 = build_demo_script()
    p2 = build_how_to_guide()
    print(f'\nAll files saved to:')
    print(f'  {p1}')
    print(f'  {p2}')
